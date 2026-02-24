"""
XLSX Preprocessing Pipeline.

Converts spreadsheet rows into one DOCX per record so documents can be ingested
by the existing doc_refresh pipeline.

Design:
- Source spreadsheets live in a preprocessing base path (per-db subfolders).
- Generated DOCX files are written under the monitored input base path.
- A state file per db_source tracks row hashes and generated paths so reruns can
  update only changed rows and remove stale generated files.

Usage:
    python -m doc_refresh.preprocessing.xlsx_to_docx --source-base <path> --output-base <path>
    python -m doc_refresh.preprocessing.xlsx_to_docx --dry-run
"""

from __future__ import annotations

import argparse
import hashlib
import io
import json
import logging
import os
import re
import tempfile
from dataclasses import dataclass
from datetime import date, datetime, time, timezone
from pathlib import Path
from typing import Any, Dict, Iterable, List, Optional, Tuple

from ..connections.file_source import FileSource, LocalFileSource, NASFileSource
from ..utils.env_config import config
from ..utils.logging_format import configure_root_logger

logger = logging.getLogger(__name__)

STATE_VERSION = 1
XLSX_EXTENSIONS = [".xlsx"]
SKIP_PREFIXES = ("~$",)
ROW_PROGRESS_INTERVAL = 5000
AUTO_KEY_CANDIDATES = {
    "id",
    "recordid",
    "record_id",
    "rowid",
    "row_id",
    "uuid",
    "guid",
    "identifier",
    "key",
    "code",
}


@dataclass
class DatabaseStats:
    """Counters for one db_source run."""

    db_source: str
    xlsx_files: int = 0
    rows_total: int = 0
    created: int = 0
    updated: int = 0
    unchanged: int = 0
    deleted: int = 0
    errors: int = 0


def _parse_args() -> argparse.Namespace:
    """Parse CLI arguments."""
    parser = argparse.ArgumentParser(
        description="Convert XLSX rows into DOCX documents for doc_refresh ingestion.",
    )
    parser.add_argument(
        "--source-base",
        type=str,
        default=os.getenv("PREPROCESSING_BASE_PATH", ""),
        help="Preprocessing source root containing <db_source> subfolders with XLSX files.",
    )
    parser.add_argument(
        "--output-base",
        type=str,
        default=os.getenv("PREPROCESSING_OUTPUT_BASE_PATH", config.BASE_PATH),
        help="Output root where generated DOCX files are written (typically doc_refresh BASE_PATH/Input).",
    )
    parser.add_argument(
        "--state-base",
        type=str,
        default=os.getenv("PREPROCESSING_STATE_PATH", ""),
        help="Path for preprocessing state files (default: <source-base>/_state).",
    )
    parser.add_argument(
        "--database-names",
        type=str,
        default=os.getenv("PREPROCESSING_DATABASE_NAMES", config.DATABASE_NAMES),
        help="Comma-separated db_source folders to process (default: auto-discover under source-base).",
    )
    parser.add_argument(
        "--file-source-mode",
        choices=["local", "nas"],
        default=config.FILE_SOURCE_MODE.lower(),
        help="Storage mode for source/output/state paths.",
    )
    parser.add_argument(
        "--generated-subdir",
        type=str,
        default="_generated/xlsx_rows",
        help="Subdirectory under each db_source where generated DOCX files are written.",
    )
    parser.add_argument(
        "--key-columns",
        type=str,
        default=os.getenv("PREPROCESSING_KEY_COLUMNS", ""),
        help="Comma-separated preferred key columns for row identity (global across sheets).",
    )
    parser.add_argument(
        "--dry-run",
        action="store_true",
        help="Plan changes without writing/deleting files or state.",
    )
    parser.add_argument(
        "--log-level",
        choices=["DEBUG", "INFO", "WARNING", "ERROR"],
        default=config.LOG_LEVEL,
        help="Logging level.",
    )
    return parser.parse_args()


def _split_csv(raw: str) -> List[str]:
    """Split comma-separated values into trimmed non-empty tokens."""
    if not raw:
        return []
    return [part.strip() for part in raw.split(",") if part.strip()]


def _canonical(value: str) -> str:
    """Lowercase, alnum-only representation for header matching."""
    return re.sub(r"[^a-z0-9]+", "", value.lower())


def _slugify(value: str, max_len: int = 80) -> str:
    """Create filesystem-safe ascii slug."""
    value = value.strip().lower()
    value = re.sub(r"[^a-z0-9]+", "_", value)
    value = value.strip("_")
    if not value:
        return ""
    return value[:max_len].rstrip("_")


def _join_remote(base: str, relative: str) -> str:
    """Join NAS share-relative paths using forward slashes."""
    base_clean = base.replace("\\", "/").strip("/")
    rel_clean = relative.replace("\\", "/").strip("/")
    if base_clean and rel_clean:
        return f"{base_clean}/{rel_clean}"
    return base_clean or rel_clean


def _storage_path(mode: str, base: str, relative: str) -> str:
    """Return absolute(local) or share-relative(nas) path."""
    if mode == "nas":
        return _join_remote(base, relative)
    return str(Path(base) / Path(relative))


def _normalize_cell(value: Any) -> str:
    """Normalize cell values to stable string content."""
    if value is None:
        return ""
    if isinstance(value, bool):
        return "true" if value else "false"
    if isinstance(value, (datetime, date, time)):
        return value.isoformat()
    if isinstance(value, float):
        if value.is_integer():
            return str(int(value))
        return f"{value:.15g}"
    return str(value).strip()


def _normalize_headers(raw_headers: Iterable[Any]) -> List[str]:
    """Create unique header names from first-row values."""
    headers: List[str] = []
    seen: Dict[str, int] = {}

    for idx, raw in enumerate(raw_headers, start=1):
        header = _normalize_cell(raw)
        if not header:
            header = f"column_{idx}"

        key = header.lower()
        seen[key] = seen.get(key, 0) + 1
        if seen[key] > 1:
            header = f"{header}_{seen[key]}"
        headers.append(header)

    return headers


def _row_hash(record: Dict[str, str]) -> str:
    """Hash normalized row record for change detection."""
    payload = json.dumps(
        record,
        sort_keys=True,
        ensure_ascii=False,
        separators=(",", ":"),
    )
    return hashlib.sha256(payload.encode("utf-8")).hexdigest()


def _rows_have_unique_nonempty(
    rows: List[Dict[str, Any]],
    columns: List[str],
) -> bool:
    """Check candidate key columns are non-empty and unique across rows."""
    if not columns:
        return False

    seen: set[Tuple[str, ...]] = set()
    for row in rows:
        record = row["record"]
        values = tuple(record.get(col, "").strip() for col in columns)
        if any(not val for val in values):
            return False
        if values in seen:
            return False
        seen.add(values)
    return True


def _resolve_key_columns(
    headers: List[str],
    rows: List[Dict[str, Any]],
    preferred: List[str],
) -> Tuple[List[str], str]:
    """Resolve key columns by preferred columns first, then auto-detection."""
    header_by_canonical = {_canonical(h): h for h in headers}

    if preferred:
        resolved: List[str] = []
        missing: List[str] = []
        for key in preferred:
            match = header_by_canonical.get(_canonical(key))
            if match:
                resolved.append(match)
            else:
                missing.append(key)
        if missing:
            logger.warning("Configured key columns not found: %s", missing)
        elif _rows_have_unique_nonempty(rows, resolved):
            return resolved, "configured"
        else:
            logger.warning(
                "Configured key columns are not unique/non-empty across rows: %s",
                resolved,
            )

    candidate_headers = [
        h for h in headers if _canonical(h) in AUTO_KEY_CANDIDATES
    ]
    for candidate in candidate_headers:
        if _rows_have_unique_nonempty(rows, [candidate]):
            return [candidate], f"auto:{candidate}"

    return [], "row_number"


def _build_row_id(
    row: Dict[str, Any],
    key_columns: List[str],
) -> str:
    """Build stable row identifier from key columns or source row number."""
    if key_columns:
        record = row["record"]
        parts = [f"{col}={record.get(col, '').strip()}" for col in key_columns]
        return "key:" + "|".join(parts)
    return f"row:{row['excel_row']}"


def _build_doc_filename(row_id: str, excel_row: int) -> str:
    """Build deterministic output filename per row."""
    display = row_id if row_id.startswith("key:") else f"row_{excel_row}"
    slug = _slugify(display, max_len=80) or f"row_{excel_row}"
    suffix = hashlib.sha1(row_id.encode("utf-8")).hexdigest()[:10]
    return f"{slug}__{suffix}.docx"


def _load_dependencies() -> Tuple[Any, Any]:
    """Load optional spreadsheet/docx dependencies with clear errors."""
    try:
        from openpyxl import load_workbook
    except ImportError as exc:
        raise RuntimeError(
            "openpyxl is required. Install with: pip install openpyxl"
        ) from exc

    try:
        from docx import Document
    except ImportError as exc:
        raise RuntimeError(
            "python-docx is required. Install with: pip install python-docx"
        ) from exc

    return load_workbook, Document


def _build_doc_bytes(
    document_cls: Any,
    db_source: str,
    workbook_relative_path: str,
    sheet_name: str,
    row: Dict[str, Any],
    key_columns: List[str],
    key_strategy: str,
) -> bytes:
    """Render one row record into a DOCX byte payload."""
    doc = document_cls()
    doc.add_heading(
        f"{db_source} / {Path(workbook_relative_path).name} / {sheet_name}",
        level=1,
    )

    meta_rows = [
        ("Record Identifier", row["row_id"]),
        ("Source Workbook", workbook_relative_path),
        ("Sheet", sheet_name),
        ("Excel Row Number", str(row["excel_row"])),
        ("Key Strategy", key_strategy),
    ]
    if key_columns:
        meta_rows.append(("Key Columns", ", ".join(key_columns)))

    for label, value in meta_rows:
        paragraph = doc.add_paragraph()
        key_run = paragraph.add_run(f"{label}: ")
        key_run.bold = True
        paragraph.add_run(value)

    doc.add_paragraph("")
    doc.add_heading("Record Fields", level=2)

    for field_name, field_value in row["record"].items():
        paragraph = doc.add_paragraph()
        key_run = paragraph.add_run(f"{field_name}: ")
        key_run.bold = True
        paragraph.add_run(field_value if field_value else "[empty]")

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def _create_source_scan_fs(mode: str, source_base: str) -> FileSource:
    """Create source scanner with base path for db folder traversal."""
    if mode == "nas":
        return NASFileSource(base_path=source_base)
    return LocalFileSource(base_path=source_base)


def _create_io_fs(mode: str) -> FileSource:
    """Create generic file IO source for absolute/share-relative paths."""
    if mode == "nas":
        return NASFileSource(base_path="")
    return LocalFileSource(base_path=None)


def _load_state(io_fs: FileSource, state_path: str) -> Dict[str, Any]:
    """Load per-db state JSON if present."""
    if not io_fs.path_exists(state_path):
        return {"version": STATE_VERSION, "rows": {}}

    with tempfile.TemporaryDirectory() as tmp:
        local_copy = io_fs.copy_to_local(state_path, tmp)
        try:
            with open(local_copy) as fh:
                payload = json.load(fh)
        except json.JSONDecodeError as exc:
            raise RuntimeError(f"State file is invalid JSON: {state_path}") from exc

    if not isinstance(payload, dict):
        raise RuntimeError(f"State file has invalid format: {state_path}")
    payload.setdefault("rows", {})
    return payload


def _write_state(io_fs: FileSource, state_path: str, state: Dict[str, Any]) -> None:
    """Write per-db state JSON."""
    data = json.dumps(state, indent=2, sort_keys=True).encode("utf-8")
    io_fs.write_data(data, state_path)


def _extract_rows_from_workbook(
    load_workbook: Any,
    local_xlsx_path: str,
    preferred_key_columns: List[str],
    workbook_label: str,
) -> List[Dict[str, Any]]:
    """Extract sheet rows from workbook with row identity and hashes."""
    workbook = load_workbook(
        filename=local_xlsx_path,
        data_only=True,
        read_only=True,
    )

    results: List[Dict[str, Any]] = []
    try:
        for worksheet in workbook.worksheets:
            logger.info(
                "[%s] Scanning sheet '%s'",
                workbook_label,
                worksheet.title,
            )
            header_row = next(
                worksheet.iter_rows(min_row=1, max_row=1, values_only=True),
                None,
            )
            if header_row is None:
                logger.info(
                    "[%s] Sheet '%s' has no header row, skipping",
                    workbook_label,
                    worksheet.title,
                )
                continue

            headers = _normalize_headers(header_row)
            rows: List[Dict[str, Any]] = []
            scanned_rows = 0

            for excel_row, values in enumerate(
                worksheet.iter_rows(min_row=2, values_only=True),
                start=2,
            ):
                scanned_rows += 1
                if scanned_rows % ROW_PROGRESS_INTERVAL == 0:
                    logger.info(
                        "[%s] Sheet '%s' scanned %d rows...",
                        workbook_label,
                        worksheet.title,
                        scanned_rows,
                    )
                record: Dict[str, str] = {}
                nonempty = False

                for idx, header in enumerate(headers):
                    cell_value = values[idx] if values and idx < len(values) else None
                    normalized = _normalize_cell(cell_value)
                    if normalized:
                        nonempty = True
                    record[header] = normalized

                if not nonempty:
                    continue

                rows.append(
                    {
                        "excel_row": excel_row,
                        "record": record,
                    }
                )

            if not rows:
                logger.info(
                    "[%s] Sheet '%s' has no non-empty rows",
                    workbook_label,
                    worksheet.title,
                )
                continue

            key_columns, key_strategy = _resolve_key_columns(
                headers=headers,
                rows=rows,
                preferred=preferred_key_columns,
            )

            for row in rows:
                row["row_id"] = _build_row_id(row, key_columns)
                row["row_hash"] = _row_hash(row["record"])

            results.append(
                {
                    "sheet_name": worksheet.title,
                    "headers": headers,
                    "key_columns": key_columns,
                    "key_strategy": key_strategy,
                    "rows": rows,
                }
            )
            logger.info(
                "[%s] Sheet '%s' extracted %d rows (key_strategy=%s, key_columns=%s)",
                workbook_label,
                worksheet.title,
                len(rows),
                key_strategy,
                key_columns or ["<row_number>"],
            )
    finally:
        workbook.close()

    return results


def _global_row_key(
    workbook_relative_path: str,
    sheet_name: str,
    row_id: str,
) -> str:
    """Build unique state key for one source row."""
    return f"{workbook_relative_path}::{sheet_name}::{row_id}"


def _build_doc_relative_path(
    generated_subdir: str,
    workbook_relative_path: str,
    sheet_name: str,
    row_id: str,
    excel_row: int,
) -> str:
    """Build generated DOCX path relative to a db_source root."""
    workbook_parts = [
        _slugify(part, max_len=80) or "part"
        for part in Path(workbook_relative_path).with_suffix("").parts
    ]
    sheet_part = _slugify(sheet_name, max_len=80) or "sheet"
    file_name = _build_doc_filename(row_id=row_id, excel_row=excel_row)

    path = Path(generated_subdir)
    for part in workbook_parts:
        path = path / part
    path = path / sheet_part / file_name
    return path.as_posix()


def _delete_if_exists(io_fs: FileSource, path: str) -> bool:
    """Delete file if present, return True when deleted."""
    if not io_fs.path_exists(path):
        return False
    io_fs.delete_file(path)
    return True


def _process_database(
    db_source: str,
    source_scan_fs: FileSource,
    io_fs: FileSource,
    mode: str,
    source_base: str,
    output_base: str,
    state_base: str,
    generated_subdir: str,
    preferred_key_columns: List[str],
    load_workbook: Any,
    document_cls: Any,
    dry_run: bool,
) -> DatabaseStats:
    """Process one db_source folder."""
    stats = DatabaseStats(db_source=db_source)

    state_rel_path = f"{db_source}.json"
    state_path = _storage_path(mode, state_base, state_rel_path)
    try:
        previous_state = _load_state(io_fs, state_path)
    except Exception as exc:
        logger.warning(
            "[%s] Failed loading state %s, continuing with empty state: %s",
            db_source,
            state_path,
            exc,
        )
        previous_state = {"version": STATE_VERSION, "rows": {}}
        stats.errors += 1

    previous_rows: Dict[str, Dict[str, Any]] = previous_state.get("rows", {})
    current_rows: Dict[str, Dict[str, Any]] = {}

    docs_to_delete: set[str] = set()

    try:
        xlsx_files = source_scan_fs.list_files(db_source, extensions=XLSX_EXTENSIONS)
    except Exception as exc:
        logger.error("Failed listing XLSX files for %s: %s", db_source, exc)
        stats.errors += 1
        return stats

    xlsx_files = [
        f for f in xlsx_files if not f.get("name", "").startswith(SKIP_PREFIXES)
    ]
    xlsx_files.sort(key=lambda item: item.get("relative_path", ""))
    stats.xlsx_files = len(xlsx_files)

    logger.info(
        "[%s] Processing %d XLSX files",
        db_source,
        stats.xlsx_files,
    )

    with tempfile.TemporaryDirectory(prefix="xlsx_preprocess_") as temp_dir:
        for file_info in xlsx_files:
            workbook_relative_path = file_info.get("relative_path", "").replace("\\", "/")
            source_file_path = file_info.get("path", "")
            logger.info("[%s] Workbook: %s", db_source, workbook_relative_path)
            try:
                local_xlsx_path = source_scan_fs.copy_to_local(source_file_path, temp_dir)
                sheets = _extract_rows_from_workbook(
                    load_workbook=load_workbook,
                    local_xlsx_path=local_xlsx_path,
                    preferred_key_columns=preferred_key_columns,
                    workbook_label=f"{db_source}/{workbook_relative_path}",
                )
            except Exception as exc:
                logger.error(
                    "[%s] Failed processing workbook %s: %s",
                    db_source,
                    workbook_relative_path,
                    exc,
                )
                stats.errors += 1
                continue

            if not sheets:
                logger.info(
                    "[%s] Workbook %s produced no record rows",
                    db_source,
                    workbook_relative_path,
                )
                continue

            workbook_rows = sum(len(sheet["rows"]) for sheet in sheets)
            logger.info(
                "[%s] Workbook %s total extracted rows: %d",
                db_source,
                workbook_relative_path,
                workbook_rows,
            )

            for sheet in sheets:
                sheet_name = sheet["sheet_name"]
                key_columns = sheet["key_columns"]
                key_strategy = sheet["key_strategy"]
                sheet_created_before = stats.created
                sheet_updated_before = stats.updated
                sheet_unchanged_before = stats.unchanged

                for row in sheet["rows"]:
                    stats.rows_total += 1
                    row_key = _global_row_key(
                        workbook_relative_path=workbook_relative_path,
                        sheet_name=sheet_name,
                        row_id=row["row_id"],
                    )
                    doc_rel_path = _build_doc_relative_path(
                        generated_subdir=generated_subdir,
                        workbook_relative_path=workbook_relative_path,
                        sheet_name=sheet_name,
                        row_id=row["row_id"],
                        excel_row=row["excel_row"],
                    )
                    doc_abs_path = _storage_path(
                        mode,
                        output_base,
                        f"{db_source}/{doc_rel_path}",
                    )

                    current_rows[row_key] = {
                        "row_hash": row["row_hash"],
                        "doc_rel_path": doc_rel_path,
                        "workbook_relative_path": workbook_relative_path,
                        "sheet_name": sheet_name,
                        "row_id": row["row_id"],
                        "excel_row": row["excel_row"],
                        "key_strategy": key_strategy,
                        "key_columns": key_columns,
                    }

                    previous = previous_rows.get(row_key)
                    if (
                        previous
                        and previous.get("row_hash") == row["row_hash"]
                        and previous.get("doc_rel_path") == doc_rel_path
                    ):
                        stats.unchanged += 1
                        continue

                    if previous and previous.get("doc_rel_path") != doc_rel_path:
                        docs_to_delete.add(previous.get("doc_rel_path", ""))

                    if dry_run:
                        if previous:
                            stats.updated += 1
                        else:
                            stats.created += 1
                        continue

                    try:
                        payload = _build_doc_bytes(
                            document_cls=document_cls,
                            db_source=db_source,
                            workbook_relative_path=workbook_relative_path,
                            sheet_name=sheet_name,
                            row=row,
                            key_columns=key_columns,
                            key_strategy=key_strategy,
                        )
                        io_fs.write_data(payload, doc_abs_path)
                        if previous:
                            stats.updated += 1
                        else:
                            stats.created += 1
                    except Exception as exc:
                        logger.error(
                            "[%s] Failed writing generated DOCX for %s (sheet=%s row=%s): %s",
                            db_source,
                            workbook_relative_path,
                            sheet_name,
                            row["excel_row"],
                            exc,
                        )
                        stats.errors += 1

                logger.info(
                    "[%s] Workbook %s sheet '%s' results: created=%d updated=%d unchanged=%d",
                    db_source,
                    workbook_relative_path,
                    sheet_name,
                    stats.created - sheet_created_before,
                    stats.updated - sheet_updated_before,
                    stats.unchanged - sheet_unchanged_before,
                )

    stale_keys = set(previous_rows.keys()) - set(current_rows.keys())
    for stale_key in stale_keys:
        stale_doc = previous_rows[stale_key].get("doc_rel_path", "")
        if stale_doc:
            docs_to_delete.add(stale_doc)

    docs_to_delete = {doc for doc in docs_to_delete if doc}
    for doc_rel_path in sorted(docs_to_delete):
        doc_abs_path = _storage_path(mode, output_base, f"{db_source}/{doc_rel_path}")
        if dry_run:
            stats.deleted += 1
            continue
        try:
            if _delete_if_exists(io_fs, doc_abs_path):
                stats.deleted += 1
        except Exception as exc:
            logger.error(
                "[%s] Failed deleting stale generated DOCX %s: %s",
                db_source,
                doc_rel_path,
                exc,
            )
            stats.errors += 1

    next_state = {
        "version": STATE_VERSION,
        "db_source": db_source,
        "source_base": source_base,
        "output_base": output_base,
        "updated_at": datetime.now(timezone.utc).isoformat(),
        "rows": current_rows,
    }

    if not dry_run:
        try:
            _write_state(io_fs, state_path, next_state)
        except Exception as exc:
            logger.error("[%s] Failed writing state file %s: %s", db_source, state_path, exc)
            stats.errors += 1

    logger.info(
        "[%s] rows=%d created=%d updated=%d unchanged=%d deleted=%d errors=%d",
        db_source,
        stats.rows_total,
        stats.created,
        stats.updated,
        stats.unchanged,
        stats.deleted,
        stats.errors,
    )
    return stats


def main() -> int:
    """CLI entrypoint."""
    args = _parse_args()
    configure_root_logger(getattr(logging, args.log_level.upper(), logging.INFO))

    source_base = args.source_base.strip()
    output_base = args.output_base.strip()
    state_base = args.state_base.strip() or _storage_path(
        args.file_source_mode,
        source_base,
        "_state",
    )

    if not source_base:
        logger.error(
            "Missing source base. Set --source-base or PREPROCESSING_BASE_PATH."
        )
        return 1
    if not output_base:
        logger.error(
            "Missing output base. Set --output-base or PREPROCESSING_OUTPUT_BASE_PATH/BASE_PATH."
        )
        return 1

    generated_subdir = args.generated_subdir.strip().strip("/")
    if not generated_subdir:
        logger.error("--generated-subdir cannot be empty")
        return 1

    preferred_key_columns = _split_csv(args.key_columns)
    db_names = _split_csv(args.database_names)

    try:
        load_workbook, document_cls = _load_dependencies()
    except RuntimeError as exc:
        logger.error("%s", exc)
        return 1

    source_scan_fs = _create_source_scan_fs(args.file_source_mode, source_base)
    io_fs = _create_io_fs(args.file_source_mode)

    if not db_names:
        try:
            db_names = source_scan_fs.list_subfolders()
        except Exception as exc:
            logger.error(
                "Failed to auto-discover db_source folders under %s: %s",
                source_base,
                exc,
            )
            return 1

    if not db_names:
        logger.warning("No db_source folders found under source base: %s", source_base)
        return 0

    logger.info("XLSX preprocessing starting")
    logger.info("  Mode: %s", args.file_source_mode)
    logger.info("  Source Base: %s", source_base)
    logger.info("  Output Base: %s", output_base)
    logger.info("  State Base: %s", state_base)
    logger.info("  Generated Subdir: %s", generated_subdir)
    logger.info("  Databases: %s", db_names)
    logger.info("  Dry Run: %s", args.dry_run)
    if preferred_key_columns:
        logger.info("  Preferred Key Columns: %s", preferred_key_columns)

    all_stats: List[DatabaseStats] = []
    for db_source in db_names:
        stats = _process_database(
            db_source=db_source,
            source_scan_fs=source_scan_fs,
            io_fs=io_fs,
            mode=args.file_source_mode,
            source_base=source_base,
            output_base=output_base,
            state_base=state_base,
            generated_subdir=generated_subdir,
            preferred_key_columns=preferred_key_columns,
            load_workbook=load_workbook,
            document_cls=document_cls,
            dry_run=args.dry_run,
        )
        all_stats.append(stats)

    total = DatabaseStats(db_source="TOTAL")
    for stats in all_stats:
        total.xlsx_files += stats.xlsx_files
        total.rows_total += stats.rows_total
        total.created += stats.created
        total.updated += stats.updated
        total.unchanged += stats.unchanged
        total.deleted += stats.deleted
        total.errors += stats.errors

    logger.info("-" * 60)
    logger.info(
        "Summary: dbs=%d xlsx=%d rows=%d created=%d updated=%d unchanged=%d deleted=%d errors=%d",
        len(all_stats),
        total.xlsx_files,
        total.rows_total,
        total.created,
        total.updated,
        total.unchanged,
        total.deleted,
        total.errors,
    )

    # Close NAS connections if used
    if hasattr(source_scan_fs, "close"):
        try:
            source_scan_fs.close()
        except Exception:
            pass
    if hasattr(io_fs, "close"):
        try:
            io_fs.close()
        except Exception:
            pass

    return 1 if total.errors > 0 else 0


if __name__ == "__main__":
    raise SystemExit(main())
